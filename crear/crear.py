#!/usr/bin/env python
from datetime import date
from datetime import timedelta
from docx import Document
from docx.shared import RGBColor
from docx.shared import Inches
from docx.shared import Cm
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import subprocess
import re
import shutil


# pip install python-docx

class Documento:
	def __init__(self):
		print('')

	# Funcion para acomodar la anchura de las celdas en la tabla
	def set_col_widths(self,table):
	    widths = (Inches(2), Inches(7), Inches(2))
	    for row in table.rows:
	        for idx, width in enumerate(widths):
	            row.cells[idx].width = width

	# Cambia los margenes de la pagina
	def set_mg_widths(self,document):
		margin = 2
		sections = document.sections
		for section in sections:
			section.top_margin = Cm(margin)
			section.bottom_margin = Cm(margin)
			section.left_margin = Cm(margin)
			section.right_margin = Cm(margin)

	def convert_to(self,folder, source, timeout=None):
		args = [self.libreoffice_exec(), '--headless', '--convert-to', 'pdf', '--outdir', folder, source]

		process = subprocess.run(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout)
		filename = re.search('-> (.*?) using filter', process.stdout.decode())

		if filename is None:
			print('error')
		else:
			return filename.group(1)


	def libreoffice_exec(self):
		# TODO: Provide support for more platforms
		if sys.platform == 'darwin':
			return '/Applications/LibreOffice.app/Contents/MacOS/soffice'
		return 'libreoffice'



	def main (self,nombre,curp,matricula,correo,tipoArchivo):
		# Se crea un nuevo documento
		document = Document()

		document.add_picture('crear/logov2.png', width=Inches(1.2))

		# Identificación del correo del alumno
		p = document.add_paragraph()
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		r = p.add_run(
		    nombre + ' <' + correo + '> \n__________________________________________________________________________________________________________________________' ).bold = True

		# Encabezado del documento
		p = document.add_paragraph()
		p.add_run('Universidad Politécnica de Victoria - Referencia Bancaria').bold = True

		# Encabezado de secretarías
		p = document.add_paragraph(
		    'Gobierno del Estado de Tamaulipas\nSecretaría de Finanzas\nSubsecretaría de Ingresos\nDirección de Recaudación\nRFC:SFG210216AJ9'
		).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		# Encabezado de boleta simple
		p = document.add_paragraph()
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		r = p.add_run('Boleta de pago, Recaudación OPD,\nFormato para pago en Ventanilla Bancaria')
		r.font.size = Pt(10)
		r.bold = True

		# Fecha del trámite
		p = document.add_paragraph()
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		r = p.add_run('Fecha Trámite: ' + str(date.today())).bold = True

		# Fecha de limite de pago
		p = document.add_paragraph()
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		r = p.add_run('\nFormato Válido hasta: ' + str(date.today() + timedelta(days=3)))
		r.font.size = Pt(10)
		r.bold = True
		font = r.font
		font.color.rgb = RGBColor(255,0,0)

		# Descripciones e informacion del pago
		document.add_paragraph(
		    'UNIVERSIDAD POLITÉCNICA DE VICTORIA\nNombre: ' + nombre + '\nCurp: ' + curp + '     Matricula: ' + matricula + '     Referencia1: 0 Referencia2: 0'
		)

		if(tipoArchivo == 'historial'):
			records = (
				(0, '------------------------------------------------------------', '---'),
			    (1, '4184 Servicio Prestado por Organismo Público Descentralizado', '50')
			)
		elif(tipoArchivo == 'credencial'):
			records = (
				(0, '------------------------------------------------------------', '---'),
			    (1, '9884 Servicio Prestado por Organismo Público Descentralizado', '120')
			)
		elif(tipoArchivo == 'kardex'):
			records = (
				(0, '------------------------------------------------------------', '---'),
			    (1, '7684 Servicio Prestado por Organismo Público Descentralizado', '180')
			)
		elif(tipoArchivo == 'constancia'):
			records = (
				(0, '------------------------------------------------------------', '---'),
			    (1, '4567 Servicio Prestado por Organismo Público Descentralizado', '80')
			)
		elif(tipoArchivo == 'toefl'):
			records = (
				(0, '------------------------------------------------------------', '---'),
			    (1, '4134 Servicio Prestado por Organismo Público Descentralizado', '800')
			)

		# Tabla que muestra la información acerca de las referencias a pagar
		table = document.add_table(rows=1, cols=3)
		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = 'Cantidad'
		hdr_cells[1].text = 'Descripcion'
		hdr_cells[2].text = 'Importe'

		total = 0
		for qty, desc, imp in records:
			if(imp == '---'):
				continue
			row_cells = table.add_row().cells
			row_cells[0].text = str(qty)
			row_cells[1].text = desc
			row_cells[2].text = '$' + imp
			total = total + float(imp)


		# Total a pagar de la referencia
		p = document.add_paragraph()
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		r = p.add_run('\nTotal a Pagar: $' + str(total)).bold = True

		# Linea de captura para el banco
		p = document.add_paragraph()
		p.add_run('\nLínea de captura Bancaria. B:3989010245446001630228019200000083387324054234\nBANAMEX').bold = True

		# Limites de la boleta e informacion extra
		p = document.add_paragraph()
		r = p.add_run('Imprimir en dos (2) tantos el formato; uno para el pagador y otro para la dependencia\nEste recibo sólo será válido cuando figure en él la certificación de nuestro sistema, sello y firma del cajero.\nPara validación de Pago o aclaraciones fiscales comunicarse al 01-800-710-65-84')
		r.font.size = Pt(8)

		# Firma de dependencia
		document.add_paragraph(
		    '\n___________________________________________________________\nFirma del contribuyente o Representante legal'
		).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT


		# Anchura de las celdas
		self.set_col_widths(table)
		self.set_mg_widths(document)

		document.save(tipoArchivo+'.docx')
		self.convert_to('archivos', tipoArchivo+'.docx')

		shutil.move(tipoArchivo+'.docx', 'archivos/'+tipoArchivo+'.docx')
		return
